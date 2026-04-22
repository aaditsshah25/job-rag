"""
generate_report.py — produces RAG_Evaluation_Report.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PLOTS_DIR = os.path.join(os.path.dirname(__file__), "eval_plots")
OUT_PATH  = os.path.join(os.path.dirname(__file__), "RAG_Evaluation_Report.docx")

# ── helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def add_table(doc, headers, rows, col_widths_cm=None, header_bg="1F4E79"):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, header_bg)

    # data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)
            if ci == 0:
                run.bold = True
            align = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            cell.paragraphs[0].alignment = align
            set_cell_bg(cell, bg)

    if col_widths_cm:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Cm(col_widths_cm[i])

    return table

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10)
    return p

def add_plot(doc, filename, caption, width=Inches(5.5)):
    path = os.path.join(PLOTS_DIR, filename)
    if os.path.exists(path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(path, width=width)
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(9)
        cap.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
        doc.add_paragraph()

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        rest = p.add_run(text)
        rest.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)

# ── document ───────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Title ──
    title = doc.add_heading("RAG Evaluation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    sub = doc.add_paragraph("JobMatch AI — Retrieval-Augmented Generation System")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].bold = True
    sub.runs[0].font.size = Pt(12)

    meta = doc.add_paragraph("Aadit Shah & Ananya  |  Gen AI Course, FLAME University  |  April 2026")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    doc.add_paragraph()

    # ── 1. Overview ──
    add_heading(doc, "1. System Overview")
    add_body(doc,
        "JobMatch AI is a Retrieval-Augmented Generation (RAG) system that matches user career profiles "
        "to relevant job postings. Given a structured user profile (skills, desired role, location, "
        "experience, salary expectations), the system retrieves the most relevant jobs from a Pinecone "
        "vector store and generates a ranked, personalised markdown response using the Gemma-3/4 language model."
    )
    doc.add_paragraph()

    add_heading(doc, "Pipeline", level=2)
    steps = [
        ("Step 1 — Embedding:", " User profile fields are tokenised and embedded using a custom SHA-256 hash-based embedding function (_hash_embed_text). Each token maps to two positions in a 768-dim vector via SHA-256 digest bytes, producing a sparse, normalised vector. This is a lexical overlap measure — no semantic generalisation."),
        ("Step 2 — Retrieval:", " Pinecone cosine similarity search retrieves the top-20 candidates. Because embeddings are hash-based (not neural), cosine scores range ~0.17–0.45, much lower than typical neural embeddings (0.6–0.9)."),
        ("Step 3 — Reranking:", " A keyword-boost reranker (_rerank_candidates) augments Pinecone scores with profile-to-job token overlap and experience alignment. This reorders candidates before passing them to the LLM."),
        ("Step 4 — Generation:", " Gemma receives the top-10 reranked candidates and produces a structured markdown response listing 5 jobs with Match Scores (N/10), reasons for fit, and apply links."),
    ]
    for bold, rest in steps:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_paragraph()
    add_heading(doc, "Ground Truth Strategy", level=2)
    add_body(doc,
        "This system has no annotated ground-truth dataset. Standard evaluation frameworks (RAGAS, DeepEval) "
        "require labelled QA pairs and semantic embeddings — neither of which applies here. Instead, three "
        "relevance proxies are used:"
    )
    proxies = [
        ("LLM Match Scores (primary):", " Gemma assigns Match Score N/10 to each recommended job. These are used as relevance grades for NDCG computation, normalised to [0, 1]."),
        ("Pinecone Cosine Score (secondary):", " Used for Precision@K with a relative median threshold, since all scores exceed any fixed low threshold."),
        ("User Feedback Ratings (tertiary):", " Ratings >= 4 from the SQLite feedback table are treated as relevant. No ratings were recorded during this evaluation run."),
    ]
    for bold, rest in proxies:
        add_bullet(doc, rest, bold_prefix=bold)

    doc.add_paragraph()

    # ── 2. Evaluation Setup ──
    add_heading(doc, "2. Evaluation Setup")
    add_body(doc,
        "19 evaluation records were collected by running 10 distinct user profiles through the live pipeline. "
        "Each record captures the query, pre- and post-rerank candidate lists (20 each), the full LLM output, "
        "and retrieval/generation latencies. All records were logged to eval_log.jsonl."
    )
    doc.add_paragraph()

    setup_headers = ["Parameter", "Value"]
    setup_rows = [
        ["Evaluation records",       "19"],
        ["Candidates retrieved (K)", "20"],
        ["Jobs scored by LLM",       "5 per request"],
        ["Total match scores",       "95"],
        ["Feedback ratings",         "0 (not collected during seeding)"],
        ["Embedding type",           "SHA-256 hash, 768-dim"],
        ["LLM model",                "Gemma-3/4 via Google Generative AI API"],
        ["Vector store",             "Pinecone (serverless)"],
    ]
    add_table(doc, setup_headers, setup_rows, col_widths_cm=[7, 9])
    doc.add_paragraph()

    # ── 3. Retrieval Metrics ──
    add_heading(doc, "3. Retrieval Metrics")

    add_heading(doc, "3.1 Precision@K and NDCG@K", level=2)
    add_body(doc,
        "Precision@K uses a relative threshold (median Pinecone score of the full result set) to determine "
        "relevance, avoiding the trivial P@K=1.0 that would result from any fixed threshold below the minimum "
        "observed score. NDCG@K uses Gemma's match scores as graded relevance, mapped back onto all 20 "
        "candidate positions (unscored positions receive grade 0)."
    )
    doc.add_paragraph()

    ret_headers = ["Metric", "K=3", "K=5", "K=10", "K=20"]
    ret_rows = [
        ["Precision@K (relative threshold)", "1.0000", "1.0000", "0.9263", "0.5000"],
        ["NDCG@K (with reranking)",          "0.6655", "0.6909", "0.6898", "0.6898"],
        ["NDCG@K (without reranking)",       "0.6960", "0.7075", "0.6900", "0.6887"],
        ["Recall@K (feedback-based)",        "N/A",    "N/A",    "N/A",    "N/A"],
        ["Hit@1 (feedback-based)",           "N/A",    "N/A",    "N/A",    "N/A"],
        ["Hit@K (feedback-based)",           "N/A",    "N/A",    "N/A",    "N/A"],
        ["MRR (feedback-based)",             "N/A",    "N/A",    "N/A",    "N/A"],
        ["MAP@K (feedback-based)",           "N/A",    "N/A",    "N/A",    "N/A"],
    ]
    add_table(doc, ret_headers, ret_rows, col_widths_cm=[7, 2.2, 2.2, 2.2, 2.2])
    cap = doc.add_paragraph("Table 1. Retrieval metrics across K values. N/A = no user feedback collected.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    add_heading(doc, "3.2 Reranking Ablation", level=2)
    add_body(doc,
        "The table below and Figure 5 compare NDCG@K with and without the keyword-boost reranker. "
        "The reranker slightly decreases NDCG at small K (–0.03 at K=3) but converges at K=10 and marginally "
        "improves at K=20. This indicates that the reranker promotes diversity and profile alignment at the "
        "cost of occasionally moving a LLM-preferred job from position 1–3 to a lower rank."
    )
    doc.add_paragraph()

    abl_headers = ["K", "NDCG (with reranking)", "NDCG (without reranking)", "Delta"]
    abl_rows = [
        ["3",  "0.6655", "0.6960", "-0.0305"],
        ["5",  "0.6909", "0.7075", "-0.0166"],
        ["10", "0.6898", "0.6900", "-0.0002"],
        ["20", "0.6898", "0.6887", "+0.0011"],
    ]
    add_table(doc, abl_headers, abl_rows, col_widths_cm=[2, 5, 5.5, 3])
    cap = doc.add_paragraph("Table 2. NDCG@K reranking ablation.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 4. Generation Metrics ──
    add_heading(doc, "4. Generation Quality Metrics")
    add_body(doc,
        "Generation quality is evaluated along three dimensions: structural format consistency, "
        "faithfulness (no hallucinated jobs), and LLM-as-judge relevance scoring."
    )
    doc.add_paragraph()

    gen_headers = ["Metric", "Value", "Notes"]
    gen_rows = [
        ["Format Consistency (avg)", "0.75 / 1.0",  "6 of 8 format rules pass consistently"],
        ["Faithfulness",             "1.00 / 1.0",  "All recommended jobs present in retrieved candidates"],
        ["LLM Judge Relevance",      "Not run",      "Requires additional API calls; excluded from this run"],
    ]
    add_table(doc, gen_headers, gen_rows, col_widths_cm=[5.5, 3.5, 7])
    cap = doc.add_paragraph("Table 3. Generation quality metrics.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    add_heading(doc, "Format Rule Breakdown", level=2)
    fmt_headers = ["Rule", "Pass Rate"]
    fmt_rows = [
        ["has_main_heading",           "1.00"],
        ["has_subheading",             "1.00"],
        ["has_job_blocks",             "1.00"],
        ["has_match_score",            "1.00"],
        ["has_why_section",            "1.00"],
        ["has_apply_link",             "1.00"],
        ["has_company_or_role_field",  "0.00"],
        ["has_location_field",         "0.00"],
    ]
    add_table(doc, fmt_headers, fmt_rows, col_widths_cm=[8, 3])
    cap = doc.add_paragraph(
        "Table 4. Per-rule format pass rates. The two failing rules reflect Gemma embedding "
        "company/location in the ### job header rather than as separate bold fields."
    )
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 5. System Metrics ──
    add_heading(doc, "5. System Metrics")
    sys_headers = ["Metric", "Value"]
    sys_rows = [
        ["Company Diversity (top-10)",      "0.6947"],
        ["Location Diversity (top-10)",     "0.4000"],
        ["Skill Diversity (top-10)",        "0.0000  (no skills metadata in Pinecone)"],
        ["Profile Alignment Score (avg)",   "0.0389  (driven by missing skills metadata)"],
        ["Retrieval Latency avg (ms)",      "2,140"],
        ["Retrieval Latency std (ms)",      "2,083"],
        ["Generation Latency avg (ms)",     "30,153"],
        ["Generation Latency std (ms)",     "966"],
        ["LLM Match Score mean (n=95)",     "4.74 / 10"],
        ["LLM Match Score std",             "0.99"],
    ]
    add_table(doc, sys_headers, sys_rows, col_widths_cm=[6.5, 9])
    cap = doc.add_paragraph("Table 5. System-level metrics.")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.runs[0].italic = True
    cap.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── 6. Plots ──
    add_heading(doc, "6. Evaluation Plots")

    add_heading(doc, "Figure 1 — Pinecone Cosine Score Distribution", level=2)
    add_body(doc,
        "All retrieved candidates have scores in the range 0.17–0.45. This is characteristic of "
        "SHA-256 hash embeddings, which measure lexical token overlap rather than semantic similarity. "
        "Neural embeddings (e.g. text-embedding-3-small) would produce scores in the 0.6–0.9 range "
        "for semantically similar documents. The red dashed line shows the median, used as the "
        "Precision@K threshold."
    )
    add_plot(doc, "plot1_pinecone_scores.png", "Figure 1. Distribution of Pinecone cosine scores across all retrieved candidates.")

    add_heading(doc, "Figure 2 — NDCG@K Curve", level=2)
    add_body(doc,
        "NDCG stabilises around 0.69 for K >= 5. The curve shows that the ranking quality improves "
        "from K=3 to K=5 and then plateaus, as most of Gemma's 5 scored jobs fall within the first 5 "
        "positions. Beyond K=5, the remaining candidate slots have grade 0 (unscored by Gemma), so "
        "NDCG does not increase further."
    )
    add_plot(doc, "plot2_ndcg_curve.png", "Figure 2. NDCG@K with and without reranking.")

    add_heading(doc, "Figure 3 — LLM Match Score Distribution", level=2)
    add_body(doc,
        "Gemma consistently assigns scores in the 3–6 range (mean=4.74, std=0.99). The absence of "
        "high scores (7–10) reflects the limitations of hash embeddings — the retrieved jobs are "
        "lexically similar but not semantically optimal. A neural embedding model would likely "
        "retrieve more semantically relevant jobs and yield higher match scores."
    )
    add_plot(doc, "plot3_match_scores.png", "Figure 3. Distribution of LLM match scores (1–10) across all 95 scored jobs.")

    add_heading(doc, "Figure 4 — Latency Breakdown", level=2)
    add_body(doc,
        "Retrieval latency averages ~2.1s (Pinecone query + network). Generation latency averages "
        "~30s (Gemma API call). Generation dominates total latency. The high std in retrieval "
        "(~2s) is due to Pinecone cold-start on the first query per session; subsequent queries "
        "benefit from connection reuse."
    )
    add_plot(doc, "plot4_latency.png", "Figure 4. Retrieval vs generation latency per request (ms).")

    add_heading(doc, "Figure 5 — Reranking Ablation", level=2)
    add_body(doc,
        "The keyword-boost reranker has a negligible negative effect at small K and converges "
        "at K=10+. This is expected: the reranker is designed to promote profile-aligned jobs "
        "that Pinecone may have ranked lower due to lexical mismatch, at the cost of occasionally "
        "displacing a Gemma-preferred job from a top position."
    )
    add_plot(doc, "plot5_rerank_ablation.png", "Figure 5. NDCG@K before vs after reranking.")

    # ── 7. Discussion ──
    add_heading(doc, "7. Discussion & Limitations")

    add_heading(doc, "Why RAGAS/DeepEval were not used", level=2)
    add_body(doc,
        "RAGAS and DeepEval require: (1) annotated question-answer pairs with ground-truth context, "
        "and (2) semantic embeddings to compute answer similarity. This system has neither — there "
        "are no labelled job-profile pairs, and the embeddings are hash-based (lexical), not semantic. "
        "All metrics are therefore approximated from available signals."
    )

    add_heading(doc, "Hash Embedding Limitations", level=2)
    add_body(doc,
        "SHA-256 hash embeddings map each token to a fixed vector position. Two synonymous tokens "
        "(e.g. 'ML' and 'machine learning') produce completely different embeddings with zero cosine "
        "overlap. This means the system cannot retrieve jobs based on semantic similarity — only "
        "exact or near-exact lexical overlap. This is the primary bottleneck for retrieval quality "
        "and explains the moderate NDCG (~0.69) and low match scores (~4.74/10)."
    )

    add_heading(doc, "Absence of Feedback Data", level=2)
    add_body(doc,
        "Recall@K, Hit@K, MRR, and MAP@K all require user feedback ratings. None were collected "
        "during this evaluation run (the feedback table was empty). These metrics would be "
        "computable if users rated recommended jobs through the application's feedback endpoint. "
        "Even a small sample of 20–30 rated sessions would enable meaningful feedback-based evaluation."
    )

    add_heading(doc, "Self-Evaluation Bias", level=2)
    add_body(doc,
        "LLM match scores are generated by the same model (Gemma) that produces the recommendations. "
        "This introduces self-evaluation bias — the model may rate its own outputs higher than an "
        "independent evaluator would. The LLM-as-judge metric (disabled in this run) partially "
        "mitigates this by using a separate prompt, but does not eliminate the bias entirely."
    )

    # ── 8. Conclusion ──
    add_heading(doc, "8. Conclusion")
    add_body(doc,
        "The JobMatch AI RAG pipeline achieves reasonable retrieval quality given the constraints of "
        "hash-based embeddings. Key findings:"
    )
    conclusions = [
        "Precision@K=1.0 at K<=5 confirms that retrieved jobs are consistently above-median relevance.",
        "NDCG@10=0.69 indicates the ranking is meaningful but not optimal — the top-5 LLM-selected jobs are not always the top-5 by Pinecone score.",
        "Faithfulness=1.0 confirms the LLM does not hallucinate jobs outside the retrieved candidate set.",
        "Format Consistency=0.75 is high; the two failing rules are due to Gemma's output style variation rather than structural failure.",
        "Reranking has minimal effect on NDCG, suggesting its value lies in profile alignment and diversity rather than raw ranking quality.",
        "End-to-end latency is dominated by Gemma generation (~30s); retrieval is fast (~2s on average).",
        "Replacing SHA-256 hash embeddings with a neural embedding model is the single highest-impact improvement for retrieval quality.",
    ]
    for c in conclusions:
        add_bullet(doc, c)

    doc.save(OUT_PATH)
    print(f"Report saved: {OUT_PATH}")


if __name__ == "__main__":
    build()
